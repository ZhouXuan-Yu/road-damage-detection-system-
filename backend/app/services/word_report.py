"""
Word 报告生成模块
使用 python-docx 生成可编辑的道路病害检测分析报告
"""

import os
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models.detection import DetectionRecord
from app.core.constants import DAMAGE_NAME_CN


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False, font_size: int = 10):
    """设置单元格文本样式"""
    cell.text = text
    para = cell.paragraphs[0]

    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 在 run 上设置字体
    for run in para.runs:
        run.font.size = Pt(font_size)
        run.font.name = "SimHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        if bold:
            run.bold = True


def set_row_background(row, color: str):
    """设置行背景色"""
    for cell in row.cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)


def generate_word_report(
    records: list[DetectionRecord],
    summary: dict,
    ai_analysis: str,
    output_path: Path,
) -> str:
    """
    生成 Word 报告

    Args:
        records: 检测记录列表
        summary: 汇总数据
        ai_analysis: AI 分析内容
        output_path: 输出文件路径

    Returns:
        str: 报告摘要
    """
    doc = Document()

    # 设置页面边距
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)

    # ==================== 标题 ====================
    title = doc.add_heading("道路病害检测分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.name = "SimHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

    # ==================== 报告概况 ====================
    doc.add_heading("一、报告概况", level=1)

    report_id = f"REP-{datetime.now().strftime('%Y%m%d%H%M%S')}"
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("报告编号", report_id),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("包含检测记录", f"{len(records)} 条"),
        ("报告格式", "Word 文档"),
    ]

    for i, (key, value) in enumerate(info_data):
        set_cell_text(info_table.rows[i].cells[0], key, bold=True)
        set_cell_text(info_table.rows[i].cells[1], value)

    # 设置表头背景色
    set_row_background(info_table.rows[0], "f3f4f6")

    doc.add_paragraph()

    # ==================== 检测概况 ====================
    doc.add_heading("二、检测概况", level=1)

    summary_table = doc.add_table(rows=3, cols=2)
    summary_table.style = "Table Grid"

    summary_data = [
        ("总检测记录", f"{summary['total_records']} 条"),
        ("病害目标总数", f"{summary['total_objects']} 个"),
        ("平均置信度", f"{summary['avg_confidence']:.2%}"),
    ]

    for i, (key, value) in enumerate(summary_data):
        set_cell_text(summary_table.rows[i].cells[0], key, bold=True)
        set_cell_text(summary_table.rows[i].cells[1], value)

    set_row_background(summary_table.rows[0], "eff6ff")

    doc.add_paragraph()

    # ==================== 病害类型分布 ====================
    doc.add_heading("三、病害类型分布", level=1)

    total = summary["total_objects"]
    sorted_dist = sorted(
        summary["damage_distribution"].items(),
        key=lambda x: -x[1]
    )

    dist_table = doc.add_table(rows=len(sorted_dist) + 1, cols=5)
    dist_table.style = "Table Grid"

    # 表头
    headers = ["序号", "病害类型", "代码", "数量", "占比"]
    for j, header in enumerate(headers):
        set_cell_text(dist_table.rows[0].cells[j], header, bold=True, center=True)

    set_row_background(dist_table.rows[0], "1e40af")
    for cell in dist_table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 数据行
    for i, (cls, count) in enumerate(sorted_dist):
        pct = count / total * 100 if total > 0 else 0
        row = dist_table.rows[i + 1]
        set_cell_text(row.cells[0], str(i + 1), center=True)
        set_cell_text(row.cells[1], DAMAGE_NAME_CN.get(cls, cls))
        set_cell_text(row.cells[2], cls, center=True)
        set_cell_text(row.cells[3], str(count), center=True)
        set_cell_text(row.cells[4], f"{pct:.1f}%", center=True)

        # 交替背景色
        if i % 2 == 1:
            set_row_background(row, "f9fafb")

    doc.add_paragraph()

    # ==================== 严重程度分布 ====================
    if summary.get("severity_distribution"):
        doc.add_heading("四、严重程度分布", level=1)

        sev_table = doc.add_table(rows=len(summary["severity_distribution"]) + 1, cols=3)
        sev_table.style = "Table Grid"

        sev_headers = ["严重程度", "数量", "说明"]
        for j, header in enumerate(sev_headers):
            set_cell_text(sev_table.rows[0].cells[j], header, bold=True, center=True)

        set_row_background(sev_table.rows[0], "1e40af")
        for cell in sev_table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)

        sev_map = {"low": "轻度", "medium": "中度", "high": "高"}
        sev_desc = {
            "low": "面积占比 < 1%",
            "medium": "面积占比 1-5%",
            "high": "面积占比 > 5%"
        }

        for i, (sev, count) in enumerate(summary["severity_distribution"].items()):
            row = sev_table.rows[i + 1]
            set_cell_text(row.cells[0], sev_map.get(sev, sev), center=True)
            set_cell_text(row.cells[1], str(count), center=True)
            set_cell_text(row.cells[2], sev_desc.get(sev, ""))

            if i % 2 == 1:
                set_row_background(row, "fef3c7")

        doc.add_paragraph()

    # ==================== 检测记录详情 ====================
    doc.add_heading("五、检测记录详情", level=1)

    detail_table = doc.add_table(rows=len(records) + 1, cols=6)
    detail_table.style = "Table Grid"

    detail_headers = ["序号", "文件名", "类型", "病害数", "置信度", "时间"]
    for j, header in enumerate(detail_headers):
        set_cell_text(detail_table.rows[0].cells[j], header, bold=True, center=True)

    set_row_background(detail_table.rows[0], "1e40af")
    for cell in detail_table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    for i, r in enumerate(records):
        row = detail_table.rows[i + 1]
        set_cell_text(row.cells[0], str(i + 1), center=True)
        set_cell_text(row.cells[1], r.filename[:30])
        set_cell_text(row.cells[2], "图片" if r.file_type == "image" else "视频", center=True)
        set_cell_text(row.cells[3], str(r.total_count), center=True)
        set_cell_text(row.cells[4], f"{r.avg_confidence:.2%}", center=True)
        ts = r.created_at.strftime("%Y-%m-%d %H:%M") if r.created_at else "-"
        set_cell_text(row.cells[5], ts, center=True)

        if i % 2 == 1:
            set_row_background(row, "f0fdf4")

    doc.add_paragraph()

    # ==================== AI 分析建议 ====================
    if ai_analysis:
        doc.add_heading("六、AI 智能分析建议", level=1)

        para = doc.add_paragraph()
        run = para.add_run(ai_analysis)
        run.font.name = "SimHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        run.font.size = Pt(10.5)
        para.paragraph_format.first_line_indent = Cm(0.74)
        para.paragraph_format.line_spacing = 1.5

    # 保存文档
    doc.save(str(output_path))

    return f"Word 报告生成成功，共 {len(records)} 条检测记录"
